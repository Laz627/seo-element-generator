import streamlit as st
from openai import OpenAI
import requests
from bs4 import BeautifulSoup
import re
from urllib.parse import quote_plus
from docx import Document
from io import BytesIO
import time

# Set page title and creator info
st.set_page_config(page_title="SEO Element Generator", layout="wide")
st.title("SEO Element Generator")
st.write("Created by Brandon Lazovic")

# Instructions
st.markdown("""
## How to use this app:
1. Enter your OpenAI API key.
2. Input up to 10 target keywords (one per line).
3. Click 'Generate SEO Elements' to get recommendations.
4. Review the results and explanations.
5. Download the results as a Word document.
""")

# Function to scrape Google search results
def scrape_google_results(keyword, num_results=10):
    url = f"https://www.google.com/search?q={quote_plus(keyword)}&num={num_results}"
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"}
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.text, "html.parser")

    results = []
    for result in soup.select("div.g"):
        title = result.select_one("h3")
        if title:
            title = title.text
            link = result.select_one("a")["href"]
            snippet = result.select_one("div.VwiC3b")
            snippet = snippet.text if snippet else ""
            results.append({"title": title, "snippet": snippet})

    return results[:num_results]

# Function to summarize competitor elements
def summarize_competitor_elements(results):
    titles = [result["title"] for result in results]
    snippets = [result["snippet"] for result in results]

    summary = f"Analyzed {len(results)} competitor results.\n"
    summary += f"Average title length: {sum(len(title) for title in titles) / len(titles):.1f} characters.\n"
    summary += f"Average snippet length: {sum(len(snippet) for snippet in snippets) / len(snippets):.1f} characters.\n"

    # Count word frequencies
    word_freq = {}
    for title in titles:
        for word in re.findall(r'\w+', title.lower()):
            if len(word) > 3:  # Ignore short words
                word_freq[word] = word_freq.get(word, 0) + 1

    common_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:5]
    summary += f"Common words in titles: {', '.join([word for word, _ in common_words])}\n\n"

    summary += "Sample competitor titles:\n"
    for title in titles[:5]:
        summary += f"- {title}\n"

    summary += "\nSample competitor snippets:\n"
    for snippet in snippets[:5]:
        summary += f"- {snippet[:100]}...\n"

    return summary

# Function to generate SEO elements using GPT-4
def generate_seo_elements(keyword, competitor_summary, api_key, max_retries=3):
    client = OpenAI(api_key=api_key)

    prompt = f"""
    Generate an H1, title tag, and meta description for the keyword: "{keyword}"
    
    Requirements:
    - H1 and title tag should be 70 characters or less
    - Meta description should be 155 characters or less
    - Avoid buzzwords and branded terms
    - Include an exact match or close variation of the target keyword
    - Closely align with the competitor results provided below
    - The elements should be a summarization of common elements from the top 10 competitors
    
    Competitor analysis:
    {competitor_summary}
    
    Based on the competitor analysis, create SEO elements that are very similar to the competitors' approach, while still being unique. Focus on common phrases, structures, and themes used by competitors.
    
    Please provide your response in the following structure:
    COMPETITOR ELEMENTS SUMMARY
    1. Most Common Title Structures:
       - [Point 1]
       - [Point 2]
       ...
    2. Common Themes in Meta Descriptions:
       - [Point 1]
       - [Point 2]
       ...

    3. Frequently Used Phrases or Keywords:
       - [Phrase 1]
       - [Phrase 2]
       ...
    4. Notable Patterns in Competitor Information Presentation:
       - [Pattern 1]
       - [Pattern 2]
       ...
    SEO ELEMENTS
    H1: [Your H1]
    Explanation:
    - [Point 1]
    - [Point 2]
    ...
    Title Tag: [Your Title Tag]
    Explanation:
    - [Point 1]
    - [Point 2]
    ...
    Meta Description: [Your Meta Description]
    Explanation:
    - [Point 1]
    - [Point 2]
    ...
    """

    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an SEO expert tasked with creating optimized on-page elements that closely align with competitor trends."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            return response.choices[0].message.content
        except Exception as e:
            if attempt == max_retries - 1:
                raise
            time.sleep(2 ** attempt)  # Exponential backoff

# Function to create Word document
def create_word_document(results):
    doc = Document()
    doc.add_heading('SEO Element Generator Results', 0)

    for result in results:
        doc.add_heading(f"Keyword: {result['Keyword']}", level=1)
        doc.add_paragraph(result['SEO Elements and Competitor Summary'])
        doc.add_heading("Competitor Analysis", level=2)
        doc.add_paragraph(result['Competitor Analysis'])
        doc.add_paragraph("\n")  # Add a blank line between results

    return doc

# Input for OpenAI API key
api_key = st.text_input("Enter your OpenAI API key:", type="password")

# Input for keywords
keywords = st.text_area("Enter up to 10 target keywords (one per line):", height=200)
keyword_list = [k.strip() for k in keywords.split("\n") if k.strip()]

if st.button("Generate SEO Elements") and api_key and keyword_list:
    results = []

    for keyword in keyword_list[:10]:
        st.subheader(f"Results for: {keyword}")

        with st.spinner(f"Analyzing competitors for '{keyword}'..."):
            competitor_results = scrape_google_results(keyword)
            competitor_summary = summarize_competitor_elements(competitor_results)

        with st.spinner(f"Generating SEO elements for '{keyword}'..."):
            seo_elements = generate_seo_elements(keyword, competitor_summary, api_key)

        st.write(seo_elements)
        st.write("Competitor Analysis Summary:")
        st.write(competitor_summary)

        results.append({
            "Keyword": keyword,
            "SEO Elements and Competitor Summary": seo_elements,
            "Competitor Analysis": competitor_summary
        })

    # Create and offer Word document download
    doc = create_word_document(results)
    bio = BytesIO()
    doc.save(bio)
    st.download_button(
        label="Download results as Word Document",
        data=bio.getvalue(),
        file_name="seo_elements_results.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.write("Please enter your OpenAI API key and at least one keyword to generate SEO elements.")
